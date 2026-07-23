from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

HERE = Path(__file__).parent
SOURCE = HERE / '阶段性科研成果报告_20260723.md'
OUT = HERE / '阶段性科研成果报告_20260723.docx'

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
GRAY = RGBColor(89, 89, 89)

def set_run(run, size=11, bold=False, color=None, italic=False):
    run.font.name = 'Hiragino Sans GB'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Hiragino Sans GB')
    run._element.rPr.rFonts.set(qn('w:cs'), 'Hiragino Sans GB')
    run._element.rPr.rFonts.set(qn('w:hint'), 'eastAsia')
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_table_widths(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), '9360'); tblW.set(qn('w:type'), 'dxa')
    tblInd = OxmlElement('w:tblInd'); tblInd.set(qn('w:w'), '120'); tblInd.set(qn('w:type'), 'dxa'); tblPr.append(tblInd)
    grid = table._tbl.tblGrid
    for idx, width in enumerate(widths):
        grid.gridCol_lst[idx].set(qn('w:w'), str(width))
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in('w:tcW')
            tcW.set(qn('w:w'), str(width)); tcW.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [2200] + [int((9360-2200)/(len(rows[0])-1))] * (len(rows[0])-1)
    if len(rows[0]) == 3: widths = [2100, 3500, 3760]
    if len(rows[0]) == 5: widths = [1300, 2050, 2100, 2050, 1860]
    set_table_widths(table, widths)
    for r, values in enumerate(rows):
        for c, value in enumerate(values):
            p = table.cell(r, c).paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            run = p.add_run(value)
            set_run(run, size=9 if r else 9.5, bold=(r == 0), color=RGBColor(0,0,0))
            if r == 0:
                shade(table.cell(r,c), 'E8EEF5')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_markdown_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.33
    # Simple inline bold / code treatment.
    parts = text.split('**')
    for i, part in enumerate(parts):
        if i % 2 == 1:
            set_run(p.add_run(part), bold=True)
        else:
            chunks = part.split('`')
            for j, chunk in enumerate(chunks):
                r = p.add_run(chunk)
                set_run(r, italic=(j % 2 == 1), color=DARK if j % 2 else None)
    return p

def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85); section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85); section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35); section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles['Normal']; normal.font.name = 'Hiragino Sans GB'; normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Hiragino Sans GB'); normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7); normal.paragraph_format.line_spacing = 1.33
    for name, size, color, before, after in [('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',12,DARK,8,4)]:
        s=styles[name]; s.font.name='Hiragino Sans GB'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Hiragino Sans GB'); s.font.size=Pt(size); s.font.color.rgb=color; s.font.bold=True; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    foot = section.footer.paragraphs[0]; foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(foot.add_run('阶段性科研成果报告｜L2 Rollup 跨层 MEV 防御'), size=8.5, color=GRAY)

    lines = SOURCE.read_text(encoding='utf-8').splitlines()
    i=0
    while i < len(lines):
        line = lines[i]
        if not line.strip(): i += 1; continue
        if line.startswith('# '):
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(28); p.paragraph_format.space_after=Pt(12)
            set_run(p.add_run(line[2:]), size=20, bold=True, color=DARK); i+=1; continue
        if line.startswith('## '):
            doc.add_paragraph(line[3:], style='Heading 1'); i+=1; continue
        if line.startswith('### '):
            doc.add_paragraph(line[4:], style='Heading 2'); i+=1; continue
        if line.startswith('|'):
            raw=[]
            while i < len(lines) and lines[i].startswith('|'):
                if not set(lines[i].replace('|','').strip()) <= set('-: '):
                    raw.append([x.strip() for x in lines[i].strip('|').split('|')])
                i += 1
            if raw: add_table(doc, raw)
            continue
        if line.startswith('**报告') or line.startswith('**研究'):
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(3)
            add_markdown_paragraph(doc, line).alignment=WD_ALIGN_PARAGRAPH.CENTER; i+=1; continue
        if line.startswith('- '):
            p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.25
            set_run(p.add_run(line[2:])); i+=1; continue
        add_markdown_paragraph(doc, line)
        i += 1
    doc.core_properties.title = '围绕 L2 Rollup 跨层 MEV 防御的阶段性科研成果报告'
    doc.core_properties.author = '杨帆'
    doc.save(OUT)

if __name__ == '__main__': main()
